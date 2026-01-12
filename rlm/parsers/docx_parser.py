"""DOCX parser using python-docx."""

from pathlib import Path

from docx import Document


def parse_docx(file_path: Path) -> str:
    """
    Extract text from a DOCX file.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Extracted text content
    """
    try:
        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)
    except Exception as e:
        return f"[DOCX Parse Error: {e}]"
