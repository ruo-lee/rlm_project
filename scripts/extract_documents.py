#!/usr/bin/env python3
"""
Extract text from PDF and DOCX files in data/law_insider/ directory.
Outputs a combined text file for use with RLM.
"""

import sys
from pathlib import Path

# Add project root to sys.path
sys.path.insert(0, str(Path(__file__).parent.parent))


def extract_pdf_text(file_path: str) -> str:
    """Extract text from a PDF file using PyMuPDF."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(file_path)
        text_parts = []

        for page_num, page in enumerate(doc, 1):
            text = page.get_text()
            if text.strip():
                text_parts.append(f"[Page {page_num}]\n{text}")

        doc.close()
        return "\n\n".join(text_parts)

    except Exception as e:
        return f"[ERROR extracting PDF: {e}]"


def extract_docx_text(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(f"[TABLE] {row_text}")

        return "\n\n".join(paragraphs)

    except Exception as e:
        return f"[ERROR extracting DOCX: {e}]"


def extract_all_documents(input_dir: str, output_file: str) -> dict:
    """
    Extract text from all PDF and DOCX files in the input directory.

    Args:
        input_dir: Directory containing PDF/DOCX files
        output_file: Path to save combined text output

    Returns:
        Statistics dictionary
    """
    input_path = Path(input_dir)

    if not input_path.exists():
        print(f"Error: Directory {input_dir} does not exist.")
        sys.exit(1)

    stats = {
        "pdf_count": 0,
        "docx_count": 0,
        "total_chars": 0,
        "files_processed": [],
        "errors": [],
    }

    combined_text = []

    # Get all PDF and DOCX files
    files = sorted(input_path.glob("*"))
    pdf_files = [f for f in files if f.suffix.lower() == ".pdf"]
    docx_files = [f for f in files if f.suffix.lower() == ".docx"]

    print(f"📂 Found {len(pdf_files)} PDF files and {len(docx_files)} DOCX files")
    print("-" * 60)

    # Process DOCX files (prefer DOCX over PDF if both exist for same document)
    processed_basenames = set()

    for docx_file in docx_files:
        print(f"📄 Extracting: {docx_file.name}...")

        text = extract_docx_text(str(docx_file))

        if text.startswith("[ERROR"):
            stats["errors"].append(docx_file.name)
            print(f"   ⚠️  Error: {text}")
        else:
            # Add document separator and content
            doc_header = f"\n{'='*80}\n📋 DOCUMENT: {docx_file.stem}\n{'='*80}\n"
            combined_text.append(doc_header + text)

            stats["docx_count"] += 1
            stats["total_chars"] += len(text)
            stats["files_processed"].append(docx_file.name)
            processed_basenames.add(docx_file.stem)

            print(f"   ✅ Extracted {len(text):,} characters")

    # Process PDF files (skip if DOCX version already processed)
    for pdf_file in pdf_files:
        # Check if we already processed a DOCX version
        pdf_basename = pdf_file.stem
        if pdf_basename in processed_basenames:
            print(f"⏭️  Skipping (DOCX version used): {pdf_file.name}")
            continue

        print(f"📄 Extracting: {pdf_file.name}...")

        text = extract_pdf_text(str(pdf_file))

        if text.startswith("[ERROR"):
            stats["errors"].append(pdf_file.name)
            print(f"   ⚠️  Error: {text}")
        else:
            doc_header = f"\n{'='*80}\n📋 DOCUMENT: {pdf_file.stem}\n{'='*80}\n"
            combined_text.append(doc_header + text)

            stats["pdf_count"] += 1
            stats["total_chars"] += len(text)
            stats["files_processed"].append(pdf_file.name)

            print(f"   ✅ Extracted {len(text):,} characters")

    # Write combined output
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    final_text = "\n\n".join(combined_text)

    with open(output_path, "w", encoding="utf-8") as f:
        f.write(final_text)

    stats["total_chars"] = len(final_text)

    return stats


def main():
    """Main entry point."""
    # Default paths
    project_root = Path(__file__).parent.parent
    input_dir = project_root / "data" / "law_insider"
    output_file = project_root / "data" / "law_insider_combined.txt"

    print("=" * 60)
    print("  📚 Law Document Text Extractor")
    print("=" * 60)
    print(f"Input:  {input_dir}")
    print(f"Output: {output_file}")
    print("=" * 60)

    stats = extract_all_documents(str(input_dir), str(output_file))

    # Print summary
    print("\n" + "=" * 60)
    print("  📊 Extraction Summary")
    print("=" * 60)
    print(f"  PDF files processed:  {stats['pdf_count']}")
    print(f"  DOCX files processed: {stats['docx_count']}")
    print(f"  Total characters:     {stats['total_chars']:,}")
    print(f"  Output file:          {output_file}")

    if stats["errors"]:
        print(f"\n  ⚠️  Errors: {len(stats['errors'])} files")
        for err in stats["errors"]:
            print(f"      - {err}")

    print("\n✅ Done! You can now use '-d 3' option in main.py")


if __name__ == "__main__":
    main()
